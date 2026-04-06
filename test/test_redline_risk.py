#!/usr/bin/env python3
"""
Test suite for redline_risk.py
"""

import json
import os
import sys
import tempfile
import unittest
from pathlib import Path

# Add parent directory to path to import redline_risk
sys.path.insert(0, str(Path(__file__).parent.parent / 'skills' / 'redline-risk' / 'tools'))

import redline_risk


class TestRedlineRisk(unittest.TestCase):
    """Test redline risk functionality."""
    
    @classmethod
    def setUpClass(cls):
        """Set up test fixtures."""
        cls.fixtures_dir = Path(__file__).parent / 'fixtures'
        cls.redline_doc = cls.fixtures_dir / 'sample-redline.docx'
        cls.v1_doc = cls.fixtures_dir / 'sample-v1.docx'
        cls.v2_doc = cls.fixtures_dir / 'sample-v2.docx'
        
        # Ensure fixtures exist
        if not cls.redline_doc.exists():
            raise FileNotFoundError(f"Test fixture not found: {cls.redline_doc}")
        if not cls.v1_doc.exists():
            raise FileNotFoundError(f"Test fixture not found: {cls.v1_doc}")
        if not cls.v2_doc.exists():
            raise FileNotFoundError(f"Test fixture not found: {cls.v2_doc}")
    
    def test_setup_check(self):
        """Test that dependencies are installed."""
        result = redline_risk.setup_check()
        self.assertEqual(result, 0, "Setup check should pass")
    
    def test_extract_changes_from_tracked_changes(self):
        """Test extracting changes from a document with tracked changes."""
        changes = redline_risk.extract_tracked_changes(str(self.redline_doc))
        
        # Should find multiple changes
        self.assertGreater(len(changes), 0, "Should extract at least one change")
        
        # Check change structure
        for change in changes:
            self.assertIn('id', change)
            self.assertIn('type', change)
            self.assertIn('section', change)
            self.assertIn('author', change)
            
            # Check type is valid
            self.assertIn(change['type'], 
                         ['insertion', 'deletion', 'modification', 'move', 
                          'formatting', 'text_formatting', 'paragraph_formatting'])
        
        # Should find specific changes we injected
        change_texts = [c.get('text', '') for c in changes]
        
        # Look for our test changes
        has_commercially_reasonable = any('commercially reasonable' in text.lower() for text in change_texts)
        has_material = any('material' in text.lower() for text in change_texts)
        
        # At least some of our changes should be found
        self.assertTrue(has_commercially_reasonable or has_material,
                       "Should find at least one of the injected changes")
    
    def test_extract_changes_from_diff(self):
        """Test extracting changes by comparing two versions."""
        changes = redline_risk.extract_changes_from_diff(
            str(self.v1_doc),
            str(self.v2_doc)
        )
        
        # Should find differences between v1 and v2
        self.assertGreater(len(changes), 0, "Should extract differences")
        
        # Check change structure
        for change in changes:
            self.assertIn('id', change)
            self.assertIn('type', change)
            self.assertIn('section', change)
    
    def test_filter_and_group_changes(self):
        """Test filtering and grouping changes."""
        # Extract changes first
        changes = redline_risk.extract_tracked_changes(str(self.redline_doc))
        
        # Filter and group
        grouped = redline_risk.filter_and_group_changes(changes)
        
        # Should produce grouped changes
        self.assertGreater(len(grouped), 0, "Should produce grouped changes")
        
        # Check grouped structure
        for group in grouped:
            self.assertIn('group_id', group)
            self.assertIn('group_type', group)
            self.assertIn('raw_change_ids', group)
            self.assertIn('section', group)
            self.assertIn('is_formatting_only', group)
            self.assertIn('is_cosmetic', group)
            
            # raw_change_ids should be a list
            self.assertIsInstance(group['raw_change_ids'], list)
    
    def test_build_output_document(self):
        """Test building the output Word document."""
        # Create sample assessments
        assessments = [
            {
                'group_id': 'g1',
                'section': '3.1',
                'section_title': 'Performance Obligations',
                'before_text': 'best efforts',
                'after_text': 'commercially reasonable efforts',
                'category': 'scope_obligation',
                'impact': 0.6,
                'direction': 'against',
                'confidence': 'high',
                'explanation': 'Replaces "best efforts" with "commercially reasonable efforts." Commercially reasonable is a lower standard.',
                'conditional_note': None
            },
            {
                'group_id': 'g2',
                'section': '6.3',
                'section_title': 'Termination for Cause',
                'before_text': 'breach',
                'after_text': 'material breach',
                'category': 'remedial_enforcement',
                'impact': 0.7,
                'direction': 'against',
                'confidence': 'high',
                'explanation': 'Adds "material" qualifier before "breach," raising the threshold for termination.',
                'conditional_note': None
            },
            {
                'group_id': 'g3',
                'section': '4.1',
                'section_title': 'Confidentiality',
                'before_text': 'all Confidential Information',
                'after_text': 'Confidential Information',
                'category': 'procedural_administrative',
                'impact': 0.1,
                'direction': 'neutral',
                'confidence': 'high',
                'explanation': 'Removes "all" before "Confidential Information." Cosmetic change with no legal effect.',
                'conditional_note': None
            }
        ]
        
        # Build output document
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            output_path = tmp.name
        
        try:
            result = redline_risk.build_output_document(
                source_path=str(self.redline_doc),
                output_path=output_path,
                title='Test Redline Analysis',
                subtitle='Vendor Services Agreement - Test',
                party='Client',
                party_mapping={'Client': 'Client Corp.', 'Vendor': 'Vendor LLC'},
                assessments=assessments
            )
            
            # Should create the output file
            self.assertTrue(os.path.exists(output_path), "Output file should be created")
            
            # Should be a valid docx (can open it)
            from docx import Document
            doc = Document(output_path)
            
            # Should have content
            self.assertGreater(len(doc.paragraphs), 0, "Document should have paragraphs")
            
            # Check for key sections
            text = '\n'.join([p.text for p in doc.paragraphs])
            self.assertIn('Executive Summary', text)
            self.assertIn('Priority Review', text)
            self.assertIn('Overall Posture', text)
        
        finally:
            # Clean up
            if os.path.exists(output_path):
                os.unlink(output_path)
    
    def test_impact_label(self):
        """Test impact label conversion."""
        self.assertEqual(redline_risk._impact_label(0.8), 'HIGH')
        self.assertEqual(redline_risk._impact_label(0.5), 'MEDIUM')
        self.assertEqual(redline_risk._impact_label(0.2), 'LOW')
    
    def test_is_cosmetic(self):
        """Test cosmetic change detection."""
        # Same text, different punctuation/whitespace
        self.assertTrue(redline_risk._is_cosmetic(
            'All Confidential Information',
            'all confidential information'
        ))
        
        # Different text
        self.assertFalse(redline_risk._is_cosmetic(
            'best efforts',
            'commercially reasonable efforts'
        ))
        
        # Empty strings
        self.assertFalse(redline_risk._is_cosmetic('', 'test'))
        self.assertFalse(redline_risk._is_cosmetic('test', ''))
    
    def test_assess_posture(self):
        """Test posture assessment."""
        # Aggressive posture (mostly against)
        assessments = [
            {'direction': 'against', 'impact': 0.8, 'category': 'risk_allocation', 
             'explanation': 'Test 1.', 'section': '1'},
            {'direction': 'against', 'impact': 0.7, 'category': 'scope_obligation', 
             'explanation': 'Test 2.', 'section': '2'},
            {'direction': 'neutral', 'impact': 0.3, 'category': 'procedural_administrative', 
             'explanation': 'Test 3.', 'section': '3'},
        ]
        
        posture, explanation, concerns = redline_risk._assess_posture(assessments)
        self.assertEqual(posture, 'aggressive')
        self.assertIn('unfavorable', explanation.lower())
        self.assertGreater(len(concerns), 0)
        
        # Balanced posture
        assessments = [
            {'direction': 'against', 'impact': 0.5, 'category': 'risk_allocation', 
             'explanation': 'Test 1.', 'section': '1'},
            {'direction': 'for', 'impact': 0.5, 'category': 'scope_obligation', 
             'explanation': 'Test 2.', 'section': '2'},
            {'direction': 'neutral', 'impact': 0.3, 'category': 'procedural_administrative', 
             'explanation': 'Test 3.', 'section': '3'},
        ]
        
        posture, explanation, concerns = redline_risk._assess_posture(assessments)
        self.assertEqual(posture, 'balanced')


class TestCLI(unittest.TestCase):
    """Test command-line interface."""
    
    @classmethod
    def setUpClass(cls):
        """Set up test fixtures."""
        cls.fixtures_dir = Path(__file__).parent / 'fixtures'
        cls.redline_doc = cls.fixtures_dir / 'sample-redline.docx'
        cls.v1_doc = cls.fixtures_dir / 'sample-v1.docx'
        cls.v2_doc = cls.fixtures_dir / 'sample-v2.docx'
    
    def test_extract_changes_cli(self):
        """Test extract-changes command via CLI simulation."""
        # Simulate CLI call
        sys.argv = ['redline_risk.py', 'extract-changes', '--source', str(self.redline_doc)]
        
        # Capture output
        from io import StringIO
        old_stdout = sys.stdout
        sys.stdout = StringIO()
        
        try:
            result = redline_risk.main()
            output = sys.stdout.getvalue()
            
            # Should succeed
            self.assertEqual(result, 0)
            
            # Should output valid JSON
            data = json.loads(output)
            self.assertIsInstance(data, list)
        
        finally:
            sys.stdout = old_stdout
    
    def test_setup_check_cli(self):
        """Test setup-check command via CLI simulation."""
        sys.argv = ['redline_risk.py', 'setup-check']
        
        result = redline_risk.main()
        self.assertEqual(result, 0)


if __name__ == '__main__':
    unittest.main()
