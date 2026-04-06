#!/usr/bin/env python3
"""
Generate synthetic test fixtures for redline-risk.

Creates Word documents with tracked changes for testing.
"""

import datetime
import zipfile
import os
from io import BytesIO
from lxml import etree
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Word OOXML namespaces
NAMESPACES = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
}

# Register namespaces
for prefix, uri in NAMESPACES.items():
    etree.register_namespace(prefix, uri)


def create_base_contract(output_path):
    """Create a basic vendor services agreement."""
    doc = Document()
    
    # Title
    title = doc.add_heading('VENDOR SERVICES AGREEMENT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    doc.add_paragraph('This Vendor Services Agreement (the "Agreement") is entered into as of March 1, 2026 (the "Effective Date") by and between:')
    doc.add_paragraph()
    
    doc.add_paragraph('CLIENT CORP., a Delaware corporation ("Client"), and', style='List Bullet')
    doc.add_paragraph('VENDOR LLC, a Delaware limited liability company ("Vendor")', style='List Bullet')
    
    doc.add_paragraph()
    
    # 1. Definitions
    doc.add_heading('1. Definitions', 1)
    doc.add_paragraph('"Services" means the software development services described in Exhibit A attached hereto.')
    doc.add_paragraph('"Deliverables" means all work product, documentation, and materials provided by Vendor to Client under this Agreement.')
    doc.add_paragraph('"Confidential Information" means all non-public information disclosed by one party to the other, whether oral or written, that is designated as confidential or that reasonably should be understood to be confidential.')
    
    # 2. Scope of Services
    doc.add_heading('2. Scope of Services', 1)
    doc.add_paragraph('2.1 Services. Vendor shall provide the Services to Client in accordance with the terms and conditions of this Agreement.')
    doc.add_paragraph('2.2 Acceptance. Client shall have fifteen (15) business days to review and accept or reject each Deliverable.')
    
    # 3. Performance Obligations
    doc.add_heading('3. Performance Obligations', 1)
    doc.add_paragraph('3.1 Standard of Care. Vendor shall use best efforts to perform the Services in a professional and workmanlike manner.')
    doc.add_paragraph('3.2 Uptime. Vendor shall use best efforts to maintain 99.9% uptime for all hosted Services.')
    
    # 4. Confidentiality
    doc.add_heading('4. Confidentiality', 1)
    doc.add_paragraph('4.1 Obligation. Each party agrees to hold all Confidential Information of the other party in strict confidence and not to disclose such information to third parties without the prior written consent of the disclosing party.')
    doc.add_paragraph('4.2 Exceptions. The confidentiality obligations shall not apply to information that: (a) is or becomes publicly available through no breach of this Agreement; (b) is rightfully received from a third party without breach of any confidentiality obligation; or (c) is independently developed without use of the Confidential Information.')
    
    # 5. Limitation of Liability
    doc.add_heading('5. Limitation of Liability', 1)
    doc.add_paragraph('5.1 Cap. IN NO EVENT SHALL EITHER PARTY\'S TOTAL LIABILITY ARISING OUT OF OR RELATED TO THIS AGREEMENT EXCEED THE AMOUNTS PAID BY CLIENT TO VENDOR IN THE TWELVE (12) MONTHS PRECEDING THE EVENT GIVING RISE TO LIABILITY.')
    doc.add_paragraph('5.2 Consequential Damages. IN NO EVENT SHALL EITHER PARTY BE LIABLE FOR ANY INDIRECT, INCIDENTAL, CONSEQUENTIAL, SPECIAL, OR EXEMPLARY DAMAGES ARISING OUT OF THIS AGREEMENT.')
    
    # 6. Term and Termination
    doc.add_heading('6. Term and Termination', 1)
    doc.add_paragraph('6.1 Term. This Agreement shall commence on the Effective Date and continue for an initial term of twelve (12) months, unless earlier terminated as provided herein.')
    doc.add_paragraph('6.2 Termination for Convenience. Either party may terminate this Agreement for convenience upon sixty (60) days prior written notice to the other party.')
    doc.add_paragraph('6.3 Termination for Cause. Either party may terminate this Agreement immediately upon written notice if the other party breaches any provision of this Agreement and fails to cure such breach within thirty (30) days after receipt of written notice thereof.')
    
    doc.save(output_path)
    print(f"Created base contract: {output_path}")


def inject_tracked_changes(docx_path, output_path):
    """
    Inject tracked changes into a Word document by manipulating the OOXML.
    
    Changes to inject:
    1. "best efforts" -> "commercially reasonable efforts" in Section 3.1
    2. "best" deleted in Section 3.2
    3. "shall" -> "may" in Section 6.2
    4. "material" inserted before "breach" in Section 6.3
    5. "sixty (60)" -> "thirty (30)" in Section 6.2 (notice period shortened)
    6. Bold added to Section 3 heading (formatting only)
    7. "all" removed before "Confidential Information" in Section 4.1 (cosmetic-ish)
    """
    
    # Open the docx as a zip
    with zipfile.ZipFile(docx_path, 'r') as docx_zip:
        # Read document.xml
        doc_xml = docx_zip.read('word/document.xml')
        tree = etree.fromstring(doc_xml)
        
        # Create namespace helpers
        w = lambda tag: f'{{{NAMESPACES["w"]}}}{tag}'
        
        # Find paragraphs containing specific text and inject changes
        all_paras = tree.findall(f'.//{w("p")}')
        
        change_id = 1
        date_str = datetime.datetime.now().isoformat()
        
        for para in all_paras:
            para_text = ''.join(para.xpath('.//w:t/text()', namespaces=NAMESPACES))
            
            # Change 1: "best efforts" -> "commercially reasonable efforts" in Section 3.1
            if 'Vendor shall use best efforts to perform' in para_text:
                _inject_replacement(para, 'best', 'commercially reasonable', change_id, date_str)
                change_id += 2
            
            # Change 2: "best" deleted in Section 3.2 (uptime)
            elif 'maintain 99.9% uptime' in para_text and 'best efforts' in para_text:
                _inject_deletion(para, 'best ', change_id, date_str)
                change_id += 1
            
            # Change 3: "shall" -> "may" in Section 6.2
            elif '6.2 Termination for Convenience' in para_text or \
                 ('Either party may terminate this Agreement for convenience' in para_text and 'upon' in para_text):
                _inject_replacement(para, 'may', 'shall', change_id, date_str)
                change_id += 2
            
            # Change 4: "material" inserted before "breach" in Section 6.3
            elif 'breaches any provision of this Agreement' in para_text:
                _inject_insertion(para, 'material ', change_id, date_str, before_word='breach')
                change_id += 1
            
            # Change 5: "sixty (60)" -> "thirty (30)" in Section 6.2
            elif 'sixty (60) days prior written notice' in para_text:
                _inject_replacement(para, 'sixty (60)', 'thirty (30)', change_id, date_str)
                change_id += 2
            
            # Change 6: Bold added to Section 3 heading (formatting only)
            elif para_text.strip() == '3. Performance Obligations':
                _inject_formatting_change(para, change_id, date_str)
                change_id += 1
            
            # Change 7: "all" removed (cosmetic)
            elif 'Each party agrees to hold all Confidential Information' in para_text:
                _inject_deletion(para, 'all ', change_id, date_str)
                change_id += 1
        
        # Write modified document
        with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as output_zip:
            for item in docx_zip.infolist():
                if item.filename == 'word/document.xml':
                    # Write our modified XML
                    output_zip.writestr(item, etree.tostring(tree, xml_declaration=True, encoding='UTF-8'))
                else:
                    # Copy other files as-is
                    output_zip.writestr(item, docx_zip.read(item.filename))
    
    print(f"Created redlined contract with tracked changes: {output_path}")


def _inject_replacement(para, old_text, new_text, change_id, date_str):
    """Inject a deletion + insertion pair (replacement) into a paragraph."""
    w = lambda tag: f'{{{NAMESPACES["w"]}}}{tag}'
    
    # Find all text runs
    runs = para.findall(f'.//{w("r")}')
    
    for run in runs:
        t_elem = run.find(f'.//{w("t")}')
        if t_elem is not None and t_elem.text and old_text in t_elem.text:
            # Split the text at the old_text occurrence
            before = t_elem.text[:t_elem.text.index(old_text)]
            after = t_elem.text[t_elem.text.index(old_text) + len(old_text):]
            
            # Clear the original run
            parent = run.getparent()
            idx = list(parent).index(run)
            
            # Create new runs with tracked changes
            if before:
                # Keep before text as-is
                t_elem.text = before
            else:
                parent.remove(run)
                idx -= 1
            
            # Insert deletion element
            del_elem = etree.Element(w('del'))
            del_elem.set(w('id'), str(change_id))
            del_elem.set(w('author'), 'Vendor Counsel')
            del_elem.set(w('date'), date_str)
            
            del_run = etree.SubElement(del_elem, w('r'))
            del_t = etree.SubElement(del_run, w('t'))
            del_t.text = old_text
            
            parent.insert(idx + 1, del_elem)
            
            # Insert insertion element
            ins_elem = etree.Element(w('ins'))
            ins_elem.set(w('id'), str(change_id + 1))
            ins_elem.set(w('author'), 'Vendor Counsel')
            ins_elem.set(w('date'), date_str)
            
            ins_run = etree.SubElement(ins_elem, w('r'))
            ins_t = etree.SubElement(ins_run, w('t'))
            ins_t.text = new_text
            
            parent.insert(idx + 2, ins_elem)
            
            # Add after text if any
            if after:
                after_run = etree.Element(w('r'))
                after_t = etree.SubElement(after_run, w('t'))
                after_t.text = after
                parent.insert(idx + 3, after_run)
            
            break


def _inject_deletion(para, text_to_delete, change_id, date_str):
    """Inject a deletion into a paragraph."""
    w = lambda tag: f'{{{NAMESPACES["w"]}}}{tag}'
    
    runs = para.findall(f'.//{w("r")}')
    
    for run in runs:
        t_elem = run.find(f'.//{w("t")}')
        if t_elem is not None and t_elem.text and text_to_delete in t_elem.text:
            before = t_elem.text[:t_elem.text.index(text_to_delete)]
            after = t_elem.text[t_elem.text.index(text_to_delete) + len(text_to_delete):]
            
            parent = run.getparent()
            idx = list(parent).index(run)
            
            if before:
                t_elem.text = before
            else:
                parent.remove(run)
                idx -= 1
            
            # Insert deletion element
            del_elem = etree.Element(w('del'))
            del_elem.set(w('id'), str(change_id))
            del_elem.set(w('author'), 'Vendor Counsel')
            del_elem.set(w('date'), date_str)
            
            del_run = etree.SubElement(del_elem, w('r'))
            del_t = etree.SubElement(del_run, w('t'))
            del_t.text = text_to_delete
            
            parent.insert(idx + 1, del_elem)
            
            if after:
                after_run = etree.Element(w('r'))
                after_t = etree.SubElement(after_run, w('t'))
                after_t.text = after
                parent.insert(idx + 2, after_run)
            
            break


def _inject_insertion(para, text_to_insert, change_id, date_str, before_word=None):
    """Inject an insertion into a paragraph."""
    w = lambda tag: f'{{{NAMESPACES["w"]}}}{tag}'
    
    runs = para.findall(f'.//{w("r")}')
    
    for run in runs:
        t_elem = run.find(f'.//{w("t")}')
        if t_elem is not None and t_elem.text and before_word and before_word in t_elem.text:
            before = t_elem.text[:t_elem.text.index(before_word)]
            after = t_elem.text[t_elem.text.index(before_word):]
            
            parent = run.getparent()
            idx = list(parent).index(run)
            
            if before:
                t_elem.text = before
            else:
                parent.remove(run)
                idx -= 1
            
            # Insert insertion element
            ins_elem = etree.Element(w('ins'))
            ins_elem.set(w('id'), str(change_id))
            ins_elem.set(w('author'), 'Vendor Counsel')
            ins_elem.set(w('date'), date_str)
            
            ins_run = etree.SubElement(ins_elem, w('r'))
            ins_t = etree.SubElement(ins_run, w('t'))
            ins_t.text = text_to_insert
            
            parent.insert(idx + 1, ins_elem)
            
            if after:
                after_run = etree.Element(w('r'))
                after_t = etree.SubElement(after_run, w('t'))
                after_t.text = after
                parent.insert(idx + 2, after_run)
            
            break


def _inject_formatting_change(para, change_id, date_str):
    """Inject a formatting-only change (add bold)."""
    w = lambda tag: f'{{{NAMESPACES["w"]}}}{tag}'
    
    # Add rPrChange to the first run
    runs = para.findall(f'.//{w("r")}')
    if runs:
        run = runs[0]
        rPr = run.find(w('rPr'))
        if rPr is None:
            rPr = etree.Element(w('rPr'))
            run.insert(0, rPr)
        
        # Add rPrChange element
        rPrChange = etree.Element(w('rPrChange'))
        rPrChange.set(w('id'), str(change_id))
        rPrChange.set(w('author'), 'Vendor Counsel')
        rPrChange.set(w('date'), date_str)
        
        rPr.insert(0, rPrChange)
        
        # Add bold
        b = etree.Element(w('b'))
        rPr.append(b)


def create_v1_and_v2():
    """Create two versions of the same contract for diff testing."""
    # Create v1 (original)
    doc = Document()
    
    doc.add_heading('VENDOR SERVICES AGREEMENT', 0)
    doc.add_paragraph()
    
    doc.add_heading('1. Services', 1)
    doc.add_paragraph('Vendor shall provide software development services to Client.')
    
    doc.add_heading('2. Payment', 1)
    doc.add_paragraph('Client shall pay Vendor $10,000 per month for the Services.')
    
    doc.add_heading('3. Term', 1)
    doc.add_paragraph('This Agreement shall continue for twelve (12) months.')
    
    v1_path = 'test/fixtures/sample-v1.docx'
    doc.save(v1_path)
    print(f"Created v1: {v1_path}")
    
    # Create v2 (modified)
    doc = Document()
    
    doc.add_heading('VENDOR SERVICES AGREEMENT', 0)
    doc.add_paragraph()
    
    doc.add_heading('1. Services', 1)
    doc.add_paragraph('Vendor shall provide software development and consulting services to Client.')
    
    doc.add_heading('2. Payment', 1)
    doc.add_paragraph('Client shall pay Vendor $12,000 per month for the Services.')
    
    doc.add_heading('3. Term', 1)
    doc.add_paragraph('This Agreement shall continue for six (6) months, with automatic renewal.')
    
    doc.add_heading('4. Confidentiality', 1)
    doc.add_paragraph('Both parties agree to maintain confidentiality of all proprietary information.')
    
    v2_path = 'test/fixtures/sample-v2.docx'
    doc.save(v2_path)
    print(f"Created v2: {v2_path}")


if __name__ == '__main__':
    # Create output directory
    os.makedirs('test/fixtures', exist_ok=True)
    
    # Create base contract
    base_path = 'test/fixtures/sample-base.docx'
    create_base_contract(base_path)
    
    # Inject tracked changes
    redline_path = 'test/fixtures/sample-redline.docx'
    inject_tracked_changes(base_path, redline_path)
    
    # Create v1 and v2 for diff testing
    create_v1_and_v2()
    
    print("\nAll fixtures created successfully!")
