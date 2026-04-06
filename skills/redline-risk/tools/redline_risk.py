#!/usr/bin/env python3
"""
Redline Risk - Contract redline analysis with directional risk scoring.

Extracts tracked changes from Word documents, groups related edits, and produces
color-coded risk assessments showing which changes favor or hurt your position.

Usage (called by the Copilot CLI skill):

    python3 redline_risk.py extract-changes --source contract-redlined.docx
    python3 redline_risk.py extract-changes --before v1.docx --after v2.docx
    python3 redline_risk.py filter-changes --changes raw-changes.json
    python3 redline_risk.py build --source contract.docx --output report.docx --changes assessments.json
    python3 redline_risk.py setup-check
"""

import argparse
import datetime
import difflib
import io
import json
import os
import re
import sys
import zipfile
from collections import defaultdict
from typing import Any, Dict, List, Optional, Tuple

try:
    from lxml import etree
except ImportError:
    etree = None

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    Document = None


# Word OOXML namespaces
NAMESPACES = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
}


def _resolve_path(path_str):
    """Expand ~ and environment variables in a user-provided path."""
    return os.path.expandvars(os.path.expanduser(path_str))


def _check_core_deps():
    """Raise if core dependencies are missing."""
    missing = []
    if etree is None:
        missing.append("lxml")
    if Document is None:
        missing.append("python-docx")
    if missing:
        print(f"Missing dependencies: {', '.join(missing)}", file=sys.stderr)
        print(f"Run setup.sh or: pip3 install lxml python-docx pillow", file=sys.stderr)
        sys.exit(1)


# ---------------------------------------------------------------------------
# Extract text from DOCX
# ---------------------------------------------------------------------------

def extract_text_from_docx(docx_path):
    """Extract full text from a Word document with paragraph structure."""
    _check_core_deps()
    doc = Document(docx_path)
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append({
                'text': text,
                'style': para.style.name if para.style else 'Normal'
            })
    return paragraphs


# ---------------------------------------------------------------------------
# Extract tracked changes from DOCX using OOXML parsing
# ---------------------------------------------------------------------------

def extract_tracked_changes(docx_path):
    """
    Extract tracked changes by parsing the raw OOXML.
    python-docx silently discards all revision markup, so we use lxml.
    """
    _check_core_deps()
    
    changes = []
    change_id = 1
    
    # Open the .docx as a zip file
    with zipfile.ZipFile(docx_path, 'r') as docx_zip:
        # Read the main document XML
        try:
            doc_xml = docx_zip.read('word/document.xml')
        except KeyError:
            print("Error: No word/document.xml found in .docx file", file=sys.stderr)
            return []
        
        # Parse with lxml
        tree = etree.fromstring(doc_xml)
        
        # Extract section structure first (for context)
        sections = _extract_sections(tree)
        
        # Find all insertions (w:ins)
        for ins in tree.xpath('.//w:ins', namespaces=NAMESPACES):
            change = _extract_insertion(ins, change_id, sections, tree)
            if change:
                changes.append(change)
                change_id += 1
        
        # Find all deletions (w:del)
        for dels in tree.xpath('.//w:del', namespaces=NAMESPACES):
            change = _extract_deletion(dels, change_id, sections, tree)
            if change:
                changes.append(change)
                change_id += 1
        
        # Find all formatting changes (w:rPrChange, w:pPrChange)
        for fmt_change in tree.xpath('.//w:rPrChange | .//w:pPrChange', namespaces=NAMESPACES):
            change = _extract_formatting_change(fmt_change, change_id, sections, tree)
            if change:
                changes.append(change)
                change_id += 1
        
        # Find all move operations (w:moveFrom, w:moveTo)
        move_from_elements = tree.xpath('.//w:moveFrom', namespaces=NAMESPACES)
        move_to_elements = tree.xpath('.//w:moveTo', namespaces=NAMESPACES)
        
        # Build move pairs by ID
        move_pairs = _pair_move_operations(move_from_elements, move_to_elements, change_id, sections, tree)
        changes.extend(move_pairs)
    
    return changes


def _extract_sections(tree):
    """Extract section structure from the document for context."""
    sections = []
    current_section = None
    
    # Look for headings and section numbers
    for para in tree.xpath('.//w:p', namespaces=NAMESPACES):
        para_style = para.find('.//w:pStyle', namespaces=NAMESPACES)
        if para_style is not None:
            style_val = para_style.get(qn('w:val'))
            if style_val and ('Heading' in style_val or 'heading' in style_val):
                text = _get_element_text(para)
                # Try to extract section number
                section_match = re.match(r'^(\d+\.?\d*\.?)\s+(.+)$', text)
                if section_match:
                    section_num = section_match.group(1)
                    section_title = section_match.group(2)
                    current_section = {
                        'number': section_num,
                        'title': section_title,
                        'element': para
                    }
                    sections.append(current_section)
                elif text:
                    current_section = {
                        'number': '',
                        'title': text,
                        'element': para
                    }
                    sections.append(current_section)
    
    return sections


def _find_containing_section(element, sections, tree):
    """Find which section contains a given element."""
    if not sections:
        return None, None
    
    # Find the element's position in document order
    all_paras = tree.xpath('.//w:p', namespaces=NAMESPACES)
    containing_para = element
    while containing_para is not None and containing_para.tag != qn('w:p'):
        containing_para = containing_para.getparent()
    
    if containing_para is None:
        return None, None
    
    try:
        element_index = all_paras.index(containing_para)
    except ValueError:
        return None, None
    
    # Find the last section heading before this element
    current_section = None
    for section in sections:
        try:
            section_index = all_paras.index(section['element'])
            if section_index <= element_index:
                current_section = section
            else:
                break
        except ValueError:
            continue
    
    if current_section:
        return current_section['number'], current_section['title']
    return None, None


def _get_element_text(element):
    """Extract all text from an element and its children."""
    texts = element.xpath('.//w:t/text()', namespaces=NAMESPACES)
    return ''.join(texts)


def _get_paragraph_text(element, tree):
    """Get the full text of the paragraph containing this element."""
    para = element
    while para is not None and para.tag != qn('w:p'):
        para = para.getparent()
    
    if para is None:
        return ""
    
    return _get_element_text(para)


def _get_context(element, tree, before_chars=30, after_chars=30):
    """Get text context before and after an element."""
    para_text = _get_paragraph_text(element, tree)
    element_text = _get_element_text(element)
    
    if not para_text or not element_text:
        return "", ""
    
    # Find element text in paragraph
    idx = para_text.find(element_text)
    if idx == -1:
        return "", ""
    
    before = para_text[max(0, idx - before_chars):idx]
    after = para_text[idx + len(element_text):idx + len(element_text) + after_chars]
    
    # Trim to word boundaries
    if before and not before[0].isspace():
        space_idx = before.find(' ')
        if space_idx > 0:
            before = before[space_idx + 1:]
    
    if after and not after[-1].isspace():
        space_idx = after.rfind(' ')
        if space_idx > 0:
            after = after[:space_idx]
    
    return before, after


def _extract_revision_metadata(element):
    """Extract author, date, and ID from revision element."""
    author = element.get(qn('w:author'), '')
    date_str = element.get(qn('w:date'), '')
    rev_id = element.get(qn('w:id'), '')
    
    # Parse date
    date = None
    if date_str:
        try:
            # Word uses ISO 8601 format
            date = datetime.datetime.fromisoformat(date_str.replace('Z', '+00:00'))
        except:
            date = None
    
    return {
        'author': author,
        'date': date.isoformat() if date else None,
        'revision_id': rev_id
    }


def _check_formatting(element):
    """Check formatting properties of an element."""
    formatting = {
        'in_table': False,
        'is_heading': False,
        'is_definition_section': False,
        'is_bold': False,
        'is_all_caps': False
    }
    
    # Check if in table
    parent = element.getparent()
    while parent is not None:
        if parent.tag == qn('w:tbl'):
            formatting['in_table'] = True
            break
        parent = parent.getparent()
    
    # Check if in heading
    para = element
    while para is not None and para.tag != qn('w:p'):
        para = para.getparent()
    
    if para is not None:
        para_style = para.find('.//w:pStyle', namespaces=NAMESPACES)
        if para_style is not None:
            style_val = para_style.get(qn('w:val'))
            if style_val and 'Heading' in style_val:
                formatting['is_heading'] = True
        
        # Check if in definition section
        para_text = _get_element_text(para)
        if re.search(r'\bdefinition\b', para_text, re.IGNORECASE):
            formatting['is_definition_section'] = True
    
    # Check bold
    rPr = element.find('.//w:rPr', namespaces=NAMESPACES)
    if rPr is not None:
        b = rPr.find('.//w:b', namespaces=NAMESPACES)
        if b is not None:
            formatting['is_bold'] = True
        
        # Check caps
        caps = rPr.find('.//w:caps', namespaces=NAMESPACES)
        if caps is not None:
            formatting['is_all_caps'] = True
    
    return formatting


def _extract_insertion(ins_element, change_id, sections, tree):
    """Extract an insertion change."""
    text = _get_element_text(ins_element)
    if not text.strip():
        return None
    
    metadata = _extract_revision_metadata(ins_element)
    section_num, section_title = _find_containing_section(ins_element, sections, tree)
    context_before, context_after = _get_context(ins_element, tree)
    para_text = _get_paragraph_text(ins_element, tree)
    formatting = _check_formatting(ins_element)
    
    return {
        'id': f'c{change_id}',
        'type': 'insertion',
        'text': text,
        'context_before': context_before,
        'context_after': context_after,
        'section': section_num or '',
        'section_title': section_title or '',
        'page': 0,  # Page numbers require more complex extraction
        'paragraph_text': para_text,
        'author': metadata['author'],
        'date': metadata['date'],
        'formatting': formatting
    }


def _extract_deletion(del_element, change_id, sections, tree):
    """Extract a deletion change."""
    text = _get_element_text(del_element)
    if not text.strip():
        return None
    
    metadata = _extract_revision_metadata(del_element)
    section_num, section_title = _find_containing_section(del_element, sections, tree)
    context_before, context_after = _get_context(del_element, tree)
    para_text = _get_paragraph_text(del_element, tree)
    formatting = _check_formatting(del_element)
    
    return {
        'id': f'c{change_id}',
        'type': 'deletion',
        'text': text,
        'context_before': context_before,
        'context_after': context_after,
        'section': section_num or '',
        'section_title': section_title or '',
        'page': 0,
        'paragraph_text': para_text,
        'author': metadata['author'],
        'date': metadata['date'],
        'formatting': formatting
    }


def _extract_formatting_change(fmt_element, change_id, sections, tree):
    """Extract a formatting-only change."""
    metadata = _extract_revision_metadata(fmt_element)
    section_num, section_title = _find_containing_section(fmt_element, sections, tree)
    para_text = _get_paragraph_text(fmt_element, tree)
    
    # Determine what kind of formatting change
    change_type = 'formatting'
    if fmt_element.tag == qn('w:pPrChange'):
        change_type = 'paragraph_formatting'
    elif fmt_element.tag == qn('w:rPrChange'):
        change_type = 'text_formatting'
    
    return {
        'id': f'c{change_id}',
        'type': change_type,
        'text': '',
        'context_before': '',
        'context_after': '',
        'section': section_num or '',
        'section_title': section_title or '',
        'page': 0,
        'paragraph_text': para_text,
        'author': metadata['author'],
        'date': metadata['date'],
        'formatting': {
            'in_table': False,
            'is_heading': False,
            'is_definition_section': False,
            'is_bold': False,
            'is_all_caps': False
        }
    }


def _pair_move_operations(move_from_list, move_to_list, start_id, sections, tree):
    """Pair moveFrom and moveTo operations by revision ID."""
    moves = []
    
    # Build lookup by revision ID
    move_from_by_id = {}
    for mf in move_from_list:
        rev_id = mf.get(qn('w:id'), '')
        if rev_id:
            move_from_by_id[rev_id] = mf
    
    change_id = start_id
    
    for mt in move_to_list:
        rev_id = mt.get(qn('w:id'), '')
        if rev_id and rev_id in move_from_by_id:
            mf = move_from_by_id[rev_id]
            
            from_text = _get_element_text(mf)
            to_text = _get_element_text(mt)
            
            if not from_text.strip():
                continue
            
            metadata = _extract_revision_metadata(mt)
            from_section, from_section_title = _find_containing_section(mf, sections, tree)
            to_section, to_section_title = _find_containing_section(mt, sections, tree)
            
            move_change = {
                'id': f'c{change_id}',
                'type': 'move',
                'text': from_text,
                'from_section': from_section or '',
                'from_section_title': from_section_title or '',
                'to_section': to_section or '',
                'to_section_title': to_section_title or '',
                'section': to_section or '',
                'section_title': to_section_title or '',
                'page': 0,
                'paragraph_text': _get_paragraph_text(mt, tree),
                'author': metadata['author'],
                'date': metadata['date'],
                'formatting': _check_formatting(mt)
            }
            moves.append(move_change)
            change_id += 1
    
    return moves


# ---------------------------------------------------------------------------
# Extract changes using before/after file comparison
# ---------------------------------------------------------------------------

def extract_changes_from_diff(before_path, after_path):
    """
    Compare two Word documents and extract changes.
    Uses sentence-level diffing.
    """
    _check_core_deps()
    
    before_paras = extract_text_from_docx(before_path)
    after_paras = extract_text_from_docx(after_path)
    
    # Align sections between documents
    before_text = [p['text'] for p in before_paras]
    after_text = [p['text'] for p in after_paras]
    
    # Use difflib to find differences
    matcher = difflib.SequenceMatcher(None, before_text, after_text)
    
    changes = []
    change_id = 1
    
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == 'equal':
            continue
        elif tag == 'delete':
            # Text deleted
            for i in range(i1, i2):
                para = before_paras[i]
                section_num, section_title = _infer_section_from_text(para['text'], para['style'])
                changes.append({
                    'id': f'c{change_id}',
                    'type': 'deletion',
                    'text': para['text'],
                    'context_before': '',
                    'context_after': '',
                    'section': section_num,
                    'section_title': section_title,
                    'page': 0,
                    'paragraph_text': para['text'],
                    'author': '',
                    'date': None,
                    'formatting': {
                        'in_table': False,
                        'is_heading': 'Heading' in para['style'],
                        'is_definition_section': False,
                        'is_bold': False,
                        'is_all_caps': False
                    }
                })
                change_id += 1
        elif tag == 'insert':
            # Text inserted
            for j in range(j1, j2):
                para = after_paras[j]
                section_num, section_title = _infer_section_from_text(para['text'], para['style'])
                changes.append({
                    'id': f'c{change_id}',
                    'type': 'insertion',
                    'text': para['text'],
                    'context_before': '',
                    'context_after': '',
                    'section': section_num,
                    'section_title': section_title,
                    'page': 0,
                    'paragraph_text': para['text'],
                    'author': '',
                    'date': None,
                    'formatting': {
                        'in_table': False,
                        'is_heading': 'Heading' in para['style'],
                        'is_definition_section': False,
                        'is_bold': False,
                        'is_all_caps': False
                    }
                })
                change_id += 1
        elif tag == 'replace':
            # Text replaced
            before_text_block = '\n'.join(before_paras[i]['text'] for i in range(i1, i2))
            after_text_block = '\n'.join(after_paras[j]['text'] for j in range(j1, j2))
            
            # Get section from after version
            if j1 < len(after_paras):
                para = after_paras[j1]
                section_num, section_title = _infer_section_from_text(para['text'], para['style'])
            else:
                section_num, section_title = '', ''
            
            changes.append({
                'id': f'c{change_id}',
                'type': 'modification',
                'old_text': before_text_block,
                'new_text': after_text_block,
                'section': section_num,
                'section_title': section_title,
                'page': 0,
                'paragraph_text': after_text_block,
                'author': '',
                'date': None,
                'formatting': {
                    'in_table': False,
                    'is_heading': False,
                    'is_definition_section': False,
                    'is_bold': False,
                    'is_all_caps': False
                }
            })
            change_id += 1
    
    return changes


def _infer_section_from_text(text, style):
    """Infer section number and title from paragraph text."""
    # Look for section number pattern
    match = re.match(r'^(\d+\.?\d*\.?)\s+(.+)$', text)
    if match:
        return match.group(1), match.group(2)
    
    # Check if it's a heading
    if 'Heading' in style:
        return '', text
    
    return '', ''


# ---------------------------------------------------------------------------
# Filter and group changes
# ---------------------------------------------------------------------------

def filter_and_group_changes(changes):
    """
    Filter formatting-only changes and group related edits.
    
    Grouping logic:
    1. Proximity grouping: insert/delete in same paragraph, same author, similar timestamp
    2. Move pairing: already done in extraction
    3. Formatting isolation: separate formatting-only changes
    4. Cosmetic detection: normalized text comparison
    """
    grouped_changes = []
    group_id = 1
    
    # First, separate formatting-only changes
    substantive_changes = []
    formatting_changes = []
    
    for change in changes:
        if change['type'] in ('formatting', 'text_formatting', 'paragraph_formatting'):
            formatting_changes.append(change)
        else:
            substantive_changes.append(change)
    
    # Group substantive changes by section and proximity
    changes_by_section = defaultdict(list)
    for change in substantive_changes:
        key = (change.get('section', ''), change.get('author', ''))
        changes_by_section[key].append(change)
    
    processed_ids = set()
    
    # Look for insert/delete pairs
    for key, section_changes in changes_by_section.items():
        # Sort by date
        section_changes.sort(key=lambda c: c.get('date') or '')
        
        for i, change in enumerate(section_changes):
            if change['id'] in processed_ids:
                continue
            
            if change['type'] == 'deletion':
                # Look for matching insertion
                for j in range(i + 1, min(i + 5, len(section_changes))):
                    other = section_changes[j]
                    if other['id'] in processed_ids:
                        continue
                    
                    if other['type'] == 'insertion':
                        # Check if they're in the same paragraph
                        if change.get('paragraph_text') == other.get('paragraph_text'):
                            # Group as replacement
                            grouped_changes.append({
                                'group_id': f'g{group_id}',
                                'group_type': 'replacement',
                                'raw_change_ids': [change['id'], other['id']],
                                'section': change['section'],
                                'section_title': change['section_title'],
                                'before_text': change['text'],
                                'after_text': other['text'],
                                'is_formatting_only': False,
                                'is_cosmetic': _is_cosmetic(change['text'], other['text'])
                            })
                            group_id += 1
                            processed_ids.add(change['id'])
                            processed_ids.add(other['id'])
                            break
    
    # Add ungrouped changes as individual groups
    for change in substantive_changes:
        if change['id'] not in processed_ids:
            if change['type'] == 'modification':
                grouped_changes.append({
                    'group_id': f'g{group_id}',
                    'group_type': 'modification',
                    'raw_change_ids': [change['id']],
                    'section': change['section'],
                    'section_title': change['section_title'],
                    'before_text': change.get('old_text', ''),
                    'after_text': change.get('new_text', ''),
                    'is_formatting_only': False,
                    'is_cosmetic': _is_cosmetic(change.get('old_text', ''), change.get('new_text', ''))
                })
            elif change['type'] == 'move':
                grouped_changes.append({
                    'group_id': f'g{group_id}',
                    'group_type': 'move',
                    'raw_change_ids': [change['id']],
                    'section': change['section'],
                    'section_title': change['section_title'],
                    'from_section': change.get('from_section', ''),
                    'to_section': change.get('to_section', ''),
                    'text': change['text'],
                    'is_formatting_only': False,
                    'is_cosmetic': False
                })
            else:
                grouped_changes.append({
                    'group_id': f'g{group_id}',
                    'group_type': change['type'],
                    'raw_change_ids': [change['id']],
                    'section': change['section'],
                    'section_title': change['section_title'],
                    'before_text': change['text'] if change['type'] == 'deletion' else '',
                    'after_text': change['text'] if change['type'] == 'insertion' else '',
                    'is_formatting_only': False,
                    'is_cosmetic': False
                })
            group_id += 1
    
    return grouped_changes


def _is_cosmetic(text1, text2):
    """Check if two text strings are cosmetically different but semantically same."""
    if not text1 or not text2:
        return False
    
    # Normalize: lowercase, collapse whitespace, remove punctuation
    def normalize(text):
        text = text.lower()
        text = re.sub(r'\s+', ' ', text)
        text = re.sub(r'[^\w\s]', '', text)
        return text.strip()
    
    return normalize(text1) == normalize(text2)


# ---------------------------------------------------------------------------
# Build output document
# ---------------------------------------------------------------------------

def build_output_document(source_path, output_path, title, subtitle, party, party_mapping, assessments):
    """
    Build the color-coded Word output document.
    
    assessments is a list of dicts with:
    - group_id
    - category
    - impact (0.0-1.0)
    - direction (for/against/neutral)
    - confidence (high/medium/low)
    - explanation
    - conditional_note (optional)
    """
    _check_core_deps()
    
    # Parse party mapping if it's a string
    if isinstance(party_mapping, str):
        party_mapping = json.loads(party_mapping)
    
    # Create document
    doc = Document()
    
    # Title page
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title)
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle_para = doc.add_paragraph()
    subtitle_run = subtitle_para.add_run(subtitle)
    subtitle_run.font.size = Pt(14)
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    date_para = doc.add_paragraph()
    date_run = date_para.add_run(datetime.datetime.now().strftime('%B %d, %Y'))
    date_run.font.size = Pt(12)
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Party perspective
    party_para = doc.add_paragraph()
    party_para.add_run('Perspective: ').bold = True
    party_para.add_run(f'{party}\n')
    
    if party_mapping:
        mapping_str = ', '.join([f'{k} = {v}' for k, v in party_mapping.items()])
        party_para.add_run('Party mapping: ').bold = True
        party_para.add_run(mapping_str)
    
    doc.add_page_break()
    
    # Summary section
    doc.add_heading('Executive Summary', 1)
    
    # Count changes
    substantive_count = len([a for a in assessments if not a.get('is_formatting_only', False)])
    
    summary_para = doc.add_paragraph()
    summary_para.add_run('Total substantive changes: ').bold = True
    summary_para.add_run(str(substantive_count))
    
    doc.add_paragraph()
    
    # Priority review
    doc.add_heading('Priority Review', 2)
    doc.add_paragraph('Changes most likely to need attention (sorted by impact):')
    
    # Sort by impact
    sorted_assessments = sorted(assessments, key=lambda a: a.get('impact', 0), reverse=True)
    
    # Top 10 or all if less
    priority_assessments = sorted_assessments[:min(10, len(sorted_assessments))]
    
    for assess in priority_assessments:
        impact_label = _impact_label(assess.get('impact', 0))
        direction = assess.get('direction', 'neutral')
        category = assess.get('category', 'unknown')
        
        item = doc.add_paragraph(style='List Bullet')
        
        # Add impact badge
        item.add_run(f'[{impact_label}] ').bold = True
        
        # Add direction indicator
        if direction == 'against':
            item.add_run('⚠️ UNFAVORABLE: ')
        elif direction == 'for':
            item.add_run('✓ FAVORABLE: ')
        else:
            item.add_run('◆ NEUTRAL: ')
        
        # Add brief summary
        explanation = assess.get('explanation', '')
        brief = explanation.split('.')[0] + '.' if explanation else 'No explanation provided.'
        item.add_run(brief)
    
    doc.add_paragraph()
    
    # Posture assessment
    doc.add_heading('Overall Posture', 2)
    posture, posture_explanation, concerns = _assess_posture(assessments)
    
    posture_para = doc.add_paragraph()
    posture_para.add_run('Classification: ').bold = True
    posture_para.add_run(posture.upper())
    
    doc.add_paragraph(posture_explanation)
    
    if concerns:
        doc.add_paragraph('Top areas of concern:')
        for concern in concerns:
            doc.add_paragraph(concern, style='List Bullet')
    
    doc.add_paragraph()
    
    # Low confidence changes
    low_conf = [a for a in assessments if a.get('confidence') == 'low']
    if low_conf:
        doc.add_heading('Changes Requiring Additional Review', 2)
        doc.add_paragraph('These changes need additional context to assess accurately:')
        
        for assess in low_conf:
            item = doc.add_paragraph(style='List Bullet')
            item.add_run(f"Section {assess.get('section', 'unknown')}: ").bold = True
            item.add_run(assess.get('explanation', ''))
            if assess.get('conditional_note'):
                doc.add_paragraph(f"Note: {assess['conditional_note']}", style='List Bullet 2')
    
    doc.add_page_break()
    
    # Detailed analysis - by severity
    doc.add_heading('Detailed Analysis (By Severity)', 1)
    
    for assess in sorted_assessments:
        _add_change_detail(doc, assess)
    
    doc.add_page_break()
    
    # Detailed analysis - by document order
    doc.add_heading('Detailed Analysis (By Document Order)', 1)
    
    # Sort by section number
    def section_sort_key(a):
        section = a.get('section', '')
        if not section:
            return (999, 0)
        parts = section.split('.')
        try:
            return tuple(int(p) for p in parts if p)
        except:
            return (999, 0)
    
    doc_order_assessments = sorted(assessments, key=section_sort_key)
    
    for assess in doc_order_assessments:
        _add_change_detail(doc, assess)
    
    # Save
    output_path = _resolve_path(output_path)
    doc.save(output_path)
    
    return output_path


def _impact_label(impact):
    """Convert numeric impact to text label."""
    if impact >= 0.7:
        return 'HIGH'
    elif impact >= 0.4:
        return 'MEDIUM'
    else:
        return 'LOW'


def _assess_posture(assessments):
    """Assess overall redline posture."""
    if not assessments:
        return 'balanced', 'No changes to assess.', []
    
    against_count = len([a for a in assessments if a.get('direction') == 'against'])
    for_count = len([a for a in assessments if a.get('direction') == 'for'])
    total = len(assessments)
    
    against_pct = against_count / total if total > 0 else 0
    for_pct = for_count / total if total > 0 else 0
    
    # High-impact unfavorable changes
    high_impact_against = [a for a in assessments 
                          if a.get('direction') == 'against' and a.get('impact', 0) >= 0.7]
    
    # Determine posture
    if against_pct > 0.6:
        posture = 'aggressive'
        explanation = f'The counterparty made {against_count} unfavorable changes ({against_pct:.0%} of all changes), with {len(high_impact_against)} high-impact changes that shift risk or obligations significantly against your position.'
    elif for_pct > 0.6:
        posture = 'defensive'
        explanation = f'The counterparty made {for_count} favorable changes ({for_pct:.0%} of all changes), suggesting they are accepting your positions or making concessions.'
    elif against_pct < 0.3 and for_pct < 0.3:
        posture = 'cleanup'
        explanation = f'Most changes are neutral or administrative. This appears to be primarily a cleanup round rather than substantive negotiation.'
    else:
        posture = 'balanced'
        explanation = f'The changes are mixed: {against_count} unfavorable, {for_count} favorable, with the remainder neutral. This suggests active negotiation on both sides.'
    
    # Top concerns
    concerns = []
    high_impact_changes = sorted([a for a in assessments if a.get('impact', 0) >= 0.6],
                                 key=lambda a: a.get('impact', 0), reverse=True)
    
    for change in high_impact_changes[:3]:
        section = change.get('section', 'unknown')
        category = change.get('category', 'unknown')
        brief = change.get('explanation', '').split('.')[0] + '.'
        concerns.append(f"Section {section} ({category}): {brief}")
    
    return posture, explanation, concerns


def _add_change_detail(doc, assess):
    """Add a detailed change entry to the document."""
    section = assess.get('section', 'unknown')
    section_title = assess.get('section_title', '')
    
    # Section heading
    if section_title:
        heading_text = f"Section {section}: {section_title}"
    else:
        heading_text = f"Section {section}"
    
    doc.add_heading(heading_text, 3)
    
    # Change text with color coding
    change_para = doc.add_paragraph()
    
    # Color background based on direction
    direction = assess.get('direction', 'neutral')
    confidence = assess.get('confidence', 'high')
    
    before_text = assess.get('before_text', '')
    after_text = assess.get('after_text', '')
    
    if before_text and after_text:
        # Replacement
        if before_text:
            del_run = change_para.add_run(before_text)
            del_run.font.strike = True
            change_para.add_run(' → ')
        
        ins_run = change_para.add_run(after_text)
        _apply_color(ins_run, direction, confidence)
    elif before_text:
        # Deletion only
        del_run = change_para.add_run(before_text)
        del_run.font.strike = True
        _apply_color(del_run, direction, confidence)
    elif after_text:
        # Insertion only
        ins_run = change_para.add_run(after_text)
        _apply_color(ins_run, direction, confidence)
    
    # Add metadata
    meta_para = doc.add_paragraph()
    
    impact_label = _impact_label(assess.get('impact', 0))
    meta_para.add_run(f"Impact: {impact_label}").bold = True
    meta_para.add_run(' | ')
    
    category = assess.get('category', 'unknown').replace('_', ' ').title()
    meta_para.add_run(f"Category: {category}").italic = True
    
    if confidence == 'low':
        meta_para.add_run(' | ')
        low_conf_run = meta_para.add_run('[LOW CONFIDENCE]')
        low_conf_run.font.color.rgb = RGBColor(255, 152, 0)
    
    # Explanation
    explanation = assess.get('explanation', '')
    if explanation:
        doc.add_paragraph(explanation)
    
    # Conditional note
    conditional_note = assess.get('conditional_note')
    if conditional_note:
        note_para = doc.add_paragraph()
        note_para.add_run('Note: ').italic = True
        note_para.add_run(conditional_note)
    
    doc.add_paragraph()


def _apply_color(run, direction, confidence):
    """Apply color highlighting based on direction and confidence."""
    # Color mapping
    if direction == 'for':
        # Green
        color = RGBColor(76, 175, 80)
        bg_color = 'C8E6C9'  # Light green
    elif direction == 'against':
        # Red
        color = RGBColor(229, 57, 53)
        bg_color = 'FFCDD2'  # Light red
    else:
        # Gray
        color = RGBColor(158, 158, 158)
        bg_color = 'E0E0E0'  # Light gray
    
    # Apply highlighting (background color)
    # For low confidence, use a lighter shade
    if confidence == 'low':
        if direction == 'for':
            bg_color = 'E8F5E9'  # Very light green
        elif direction == 'against':
            bg_color = 'FFEBEE'  # Very light red
        else:
            bg_color = 'F5F5F5'  # Very light gray
    
    # Apply shading to paragraph
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), bg_color)
    run._element.get_or_add_rPr().append(shading_elm)


# ---------------------------------------------------------------------------
# Setup check
# ---------------------------------------------------------------------------

def setup_check():
    """Verify all dependencies are installed."""
    print("Checking dependencies...", file=sys.stderr)
    
    all_good = True
    
    # Check lxml
    if etree is None:
        print("❌ lxml not found", file=sys.stderr)
        all_good = False
    else:
        print("✓ lxml installed", file=sys.stderr)
    
    # Check python-docx
    if Document is None:
        print("❌ python-docx not found", file=sys.stderr)
        all_good = False
    else:
        print("✓ python-docx installed", file=sys.stderr)
    
    # Check pillow
    try:
        from PIL import Image
        print("✓ pillow installed", file=sys.stderr)
    except ImportError:
        print("❌ pillow not found", file=sys.stderr)
        all_good = False
    
    if all_good:
        print("\n✓ All dependencies installed", file=sys.stderr)
        return 0
    else:
        print("\nMissing dependencies. Run: pip3 install lxml python-docx pillow", file=sys.stderr)
        return 1


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description='Redline Risk - Contract redline analysis with directional risk scoring'
    )
    
    subparsers = parser.add_subparsers(dest='command', help='Command to run')
    
    # extract-changes
    extract_parser = subparsers.add_parser('extract-changes', help='Extract tracked changes from document')
    extract_parser.add_argument('--source', help='Path to redlined .docx file')
    extract_parser.add_argument('--before', help='Path to original version (for diff mode)')
    extract_parser.add_argument('--after', help='Path to modified version (for diff mode)')
    
    # filter-changes
    filter_parser = subparsers.add_parser('filter-changes', help='Filter and group changes')
    filter_parser.add_argument('--changes', required=True, help='Path to raw changes JSON file')
    
    # build
    build_parser = subparsers.add_parser('build', help='Build color-coded output document')
    build_parser.add_argument('--source', required=True, help='Path to source document')
    build_parser.add_argument('--output', required=True, help='Path to output .docx file')
    build_parser.add_argument('--title', required=True, help='Document title')
    build_parser.add_argument('--subtitle', required=True, help='Document subtitle')
    build_parser.add_argument('--party', required=True, help='Your party name')
    build_parser.add_argument('--party-mapping', help='JSON mapping of party names')
    build_parser.add_argument('--changes', required=True, help='JSON array of assessed changes')
    
    # setup-check
    setup_parser = subparsers.add_parser('setup-check', help='Verify dependencies')
    
    args = parser.parse_args()
    
    if not args.command:
        parser.print_help()
        return 1
    
    try:
        if args.command == 'extract-changes':
            if args.source:
                # Extract from tracked changes
                changes = extract_tracked_changes(_resolve_path(args.source))
            elif args.before and args.after:
                # Extract from diff
                changes = extract_changes_from_diff(_resolve_path(args.before), _resolve_path(args.after))
            else:
                print("Error: Must provide --source OR --before and --after", file=sys.stderr)
                return 1
            
            print(json.dumps(changes, indent=2))
            return 0
        
        elif args.command == 'filter-changes':
            # Load changes from file
            with open(_resolve_path(args.changes), 'r') as f:
                changes = json.load(f)
            
            grouped = filter_and_group_changes(changes)
            print(json.dumps(grouped, indent=2))
            return 0
        
        elif args.command == 'build':
            # Load assessments
            if args.changes.startswith('['):
                # JSON string
                assessments = json.loads(args.changes)
            else:
                # File path
                with open(_resolve_path(args.changes), 'r') as f:
                    assessments = json.load(f)
            
            output_path = build_output_document(
                _resolve_path(args.source),
                args.output,
                args.title,
                args.subtitle,
                args.party,
                args.party_mapping,
                assessments
            )
            
            print(json.dumps({'output': output_path}))
            return 0
        
        elif args.command == 'setup-check':
            return setup_check()
    
    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        import traceback
        traceback.print_exc(file=sys.stderr)
        return 1


if __name__ == '__main__':
    sys.exit(main())
